from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os

def save_resume_as_docx(content: str, filename: str):
    try:
        # Ensure the directory exists
        output_dir = "generated_resumes"
        os.makedirs(output_dir, exist_ok=True)
        
        doc = Document()
        doc.add_heading('Resume', 0)
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('#'):
                    doc.add_heading(line.strip('# '), level=1)
                else:
                    doc.add_paragraph(line)
        path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(path)
        return path
    except Exception as e:
        print(f"Error saving DOCX: {e}")
        raise

def save_resume_as_pdf(content: str, filename: str):
    try:
        # Ensure the directory exists
        output_dir = "generated_resumes"
        os.makedirs(output_dir, exist_ok=True)
        
        path = os.path.join(output_dir, f"{filename}.pdf")
        c = canvas.Canvas(path, pagesize=letter)
        y = 750
        for line in content.split('\n'):
            if line.strip():
                c.drawString(50, y, line[:100])  # truncate long lines
            y -= 18
        c.save()
        return path
    except Exception as e:
        print(f"Error saving PDF: {e}")
        raise